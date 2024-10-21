import pandas as pd
from typing import List
import os
from langchain.schema import Document

# Function to read Excel and CSV files
def read_file(file_path: str) -> pd.DataFrame:
    """
    Reads xlsx, xls, and csv files and returns a pandas DataFrame.

    Args:
        file_path (str): The file path to the Excel or CSV file.

    Returns:
        pd.DataFrame: The content of the file in a pandas DataFrame.
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == ".xlsx":
        return pd.read_excel(file_path, engine='openpyxl')
    elif file_extension == ".xls":
        return pd.read_excel(file_path, engine='xlrd')
    elif file_extension == ".csv":
        return pd.read_csv(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")

# Function to convert DataFrame content to text
def dataframe_to_text(df: pd.DataFrame) -> str:
    """
    Converts the content of a pandas DataFrame to text for use in a RAG system.

    Args:
        df (pd.DataFrame): The DataFrame to convert.

    Returns:
        str: The text representation of the DataFrame.
    """
    return df.to_string(index=False)

# Function to convert file content to LangChain Documents
def file_to_documents(file_path: str) -> List[Document]:
    """
    Converts the content of an Excel or CSV file to a list of LangChain Documents.

    Args:
        file_path (str): Path to the file.

    Returns:
        List[Document]: A list of LangChain documents containing the file content.
    """
    df = read_file(file_path)
    text = dataframe_to_text(df)
    
    # Assuming each document corresponds to a single file in this case.
    return [Document(page_content=text, metadata={"source": file_path})]

# Example usage
if __name__ == "__main__":
    file_path = "path/to/your/file.xlsx"  # Provide the path to your file
    documents = file_to_documents(file_path)

    # Now you can add these documents to your LangChain vector store for retrieval.
    for doc in documents:
        print(f"Document from {doc.metadata['source']}:\n{doc.page_content}\n")



=============
import io
import pytesseract
from PIL import Image
import pdfplumber

# Function to extract text from images using pytesseract
def extract_text_from_image(image_bytes):
    image = Image.open(io.BytesIO(image_bytes))
    text = pytesseract.image_to_string(image)
    return text

# Function to read text and images from the PDF and insert extracted image text at the correct position
def read_pdf(file_path):
    content = ""
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = ""
            # Extract the page text and sort the objects based on their Y positions (to maintain order)
            words = page.extract_words()
            words = sorted(words, key=lambda x: x['top'])  # Sorting by position to maintain reading flow

            image_objects = page.images
            image_index = 0  # To track images on the page
            
            for word in words:
                page_text += word['text'] + " "
                
                # Check if there's an image near the current word position and insert OCR text
                if image_index < len(image_objects):
                    image = image_objects[image_index]
                    if word['top'] >= image['top']:
                        image_bytes = extract_image(page, image)
                        ocr_text = extract_text_from_image(image_bytes)
                        page_text += f"[Extracted from image]: {ocr_text}\n"
                        image_index += 1  # Move to the next image

            # Add the text for this page to the content
            content += page_text + "\n"
    
    return content

# Helper function to extract an image from a page
def extract_image(page, image_obj):
    x0, top, x1, bottom = image_obj['x0'], image_obj['top'], image_obj['x1'], image_obj['bottom']
    cropped_image = page.within_bbox((x0, top, x1, bottom)).to_image()
    img_byte_arr = io.BytesIO()
    cropped_image.save(img_byte_arr, format='PNG')  # Save image as PNG in memory
    return img_byte_arr.getvalue()

if __name__ == "__main__":
    pdf_file = 'your_document.pdf'
    extracted_content = read_pdf(pdf_file)
    print(extracted_content)



===================
import docx
import easyocr
from PIL import Image
import io

# Initialize EasyOCR reader (you can specify the language(s) you need)
reader = easyocr.Reader(['en'], gpu=False)  # Use GPU if available by setting gpu=True

# Function to extract text from images using EasyOCR
def extract_text_from_image(image_bytes):
    image = Image.open(io.BytesIO(image_bytes))
    image.save('temp_image.png')  # Save image temporarily for EasyOCR to process
    result = reader.readtext('temp_image.png', detail=0)  # Extract text with EasyOCR
    return " ".join(result)

# Function to extract text and images from docx, and place text where images are
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = ""

    for para in doc.paragraphs:
        content += para.text + "\n"  # Extract the paragraph text

        # Check for images in the paragraph's inline shapes (images)
        for run in para.runs:
            if run._element.xpath('.//a:blip' or './/pic:blip'):
                # This looks for image references within the run
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        # Extract text from the image
                        ocr_text = extract_text_from_image(image_bytes)
                        # Add extracted text where the image was
                        content += f"[Extracted from image]: {ocr_text}\n"
    
    return content
    
    
import docx
import pytesseract
from PIL import Image
import io

# Function to extract text from images using pytesseract
def extract_text_from_image(image_bytes):
    image = Image.open(io.BytesIO(image_bytes))
    text = pytesseract.image_to_string(image)
    return text

# Function to extract text and images from docx, and place text where images are
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = ""

    for para in doc.paragraphs:
        content += para.text + "\n"  # Extract the paragraph text

        # Check for images in the paragraph's inline shapes (images)
        for run in para.runs:
            if run._element.xpath('.//a:blip' or './/pic:blip'):
                # This looks for image references within the run
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        # Extract text from the image
                        ocr_text = extract_text_from_image(image_bytes)
                        # Add extracted text where the image was
                        content += f"[Extracted from image]: {ocr_text}\n"
    
    return content

if __name__ == "__main__":
    docx_file = 'your_document.docx'
    extracted_content = read_docx(docx_file)
    print(extracted_content)


===========
import docx
import pytesseract
from PIL import Image
import io

# Function to extract text from images using pytesseract
def extract_text_from_image(image_bytes):
    image = Image.open(io.BytesIO(image_bytes))
    text = pytesseract.image_to_string(image)
    return text

# Function to read text and images from docx
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = ""
    
    # Loop through document elements (paragraphs, runs, and images)
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.tag.endswith('drawing'):  # This detects images
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        # Extract text from image and append where the image is
                        ocr_text = extract_text_from_image(image_bytes)
                        content += f"\n[Extracted from image]: {ocr_text}\n"
            else:
                content += run.text

        content += "\n"

    return content

if __name__ == "__main__":
    docx_file = 'your_document.docx'
    extracted_content = read_docx(docx_file)
    print(extracted_content)


==================================
from sqlalchemy.types import TypeDecorator, Text
import json

class JsonClob(TypeDecorator):
    """Custom SQLAlchemy type for storing JSON data in a CLOB column."""
    impl = Text

    def process_bind_param(self, value, dialect):
        """Convert Python dictionary to JSON string when saving to the database."""
        if value is None:
            return None
        return json.dumps(value)

    def process_result_value(self, value, dialect):
        """Convert JSON string from the database to a Python dictionary."""
        if value is None:
            return {}
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return {}



class Bucket(Base):
    __tablename__ = 'buckets'

    unique_name = Column(String(255), primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(UUID(as_uuid=True), ForeignKey('users.id'), nullable=False)
    name = Column(String(255), nullable=False)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)

    # Relationship to User
    user = relationship("User", back_populates="buckets")
    files = relationship("File", back_populates="bucket")

class File(Base):
    __tablename__ = 'files'

    unique_name = Column(String(255), primary_key=True, default=lambda: str(uuid.uuid4()))
    bucket_id = Column(String(255), ForeignKey('buckets.unique_name'), nullable=False)
    name = Column(String(255), nullable=False)
    size = Column(BigInteger, nullable=False)
    content_type = Column(String(255))
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)

    # Relationship to Bucket
    bucket = relationship("Bucket", back_populates="files")


from transformers import AutoModelForCausalLM, AutoTokenizer

def create_chatbot(model_name):
  """Creates a chatbot using a Hugging Face LLM."""
  model = AutoModelForCausalLM.from_pretrained(model_name)
  tokenizer = AutoTokenizer.from_pretrained(model_name)

  def chatbot(user_input):
    inputs = tokenizer(user_input, return_tensors="pt")
    outputs = model.generate(**inputs, max_length=100)  # Adjust max_length as needed
    response = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return response

  return chatbot



import mistral

def create_chatbot():
  """Creates a chatbot using the Mistral language model."""
  llm = mistral.load("mistral-7b-instruct")  # Replace with your desired Mistral model

  def chatbot(user_input):
    """Handles user input and generates a response."""
    response = llm.generate(user_input, max_tokens=100)  # Adjust max_tokens as needed
    return response.text

  return chatbot





Updated Database Schema
Tables
Users

user_id: Primary Key, unique identifier for each user.
username: User's name.
email: User's email address.
password: User's hashed password.
Buckets

bucket_id: Primary Key, unique identifier for each bucket or sub-bucket.
bucket_name: Name of the bucket or sub-bucket (e.g., userid/private_buckets/bkt1, sub_bucket).
parent_bucket_id: Foreign Key, references bucket_id in the same table, for linking sub-buckets to their parent buckets (NULL for top-level buckets).
owner_id: Foreign Key, references user_id in the Users table, to store the owner/creator of the bucket.
AccessControl

access_id: Primary Key, unique identifier for each access control entry.
user_id: Foreign Key, references user_id in the Users table.
bucket_id: Foreign Key, references bucket_id in the Buckets table.
access_level: Type of access granted (e.g., read, write, admin).


from sqlalchemy import create_engine, Column, Integer, String, ForeignKey, Enum
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship, sessionmaker

Base = declarative_base()

class User(Base):
    __tablename__ = 'users'
    
    user_id = Column(Integer, primary_key=True, autoincrement=True)
    username = Column(String(255), nullable=False)
    email = Column(String(255), nullable=False)
    password = Column(String(255), nullable=False)
    
    owned_buckets = relationship("Bucket", back_populates="owner")
    access_controls = relationship("AccessControl", back_populates="user")

class Bucket(Base):
    __tablename__ = 'buckets'
    
    bucket_id = Column(Integer, primary_key=True, autoincrement=True)
    bucket_name = Column(String(255), nullable=False)
    parent_bucket_id = Column(Integer, ForeignKey('buckets.bucket_id'), nullable=True)
    owner_id = Column(Integer, ForeignKey('users.user_id'), nullable=False)
    
    parent_bucket = relationship("Bucket", remote_side=[bucket_id])
    sub_buckets = relationship("Bucket", back_populates="parent_bucket", remote_side=[parent_bucket_id])
    owner = relationship("User", back_populates="owned_buckets")
    access_controls = relationship("AccessControl", back_populates="bucket")

class AccessControl(Base):
    __tablename__ = 'access_control'
    
    access_id = Column(Integer, primary_key=True, autoincrement=True)
    user_id = Column(Integer, ForeignKey('users.user_id'), nullable=False)
    bucket_id = Column(Integer, ForeignKey('buckets.bucket_id'), nullable=False)
    access_level = Column(Enum('read', 'write', 'admin', name='access_level_enum'), nullable=False)
    
    user = relationship("User", back_populates="access_controls")
    bucket = relationship("Bucket", back_populates="access_controls")

# Database connection setup
DATABASE_URL = "mysql+pymysql://user:password@localhost/dbname"  # Replace with your actual database URL

engine = create_engine(DATABASE_URL)
Session = sessionmaker(bind=engine)
session = Session()

# Create tables
Base.metadata.create_all(engine)










def get_history_user(db: Session, user: str):
    logger.info('get_history_user!')
    
    # Query the UserSession table filtered by the user and ordered by session_id and created_at
    sessions = db.query(UserSession).filter(UserSession.nbkid == user).order_by(UserSession.session_id, UserSession.created_at).all()
    
    # Group the sessions by session_id
    session_dict = defaultdict(list)
    for session in sessions:
        session_dict[session.session_id].append({
            'message_type': session.message_type,
            'message': session.message,
            'created_at': session.created_at
        })
    
    # Sort messages within each session by created_at in ascending order
    for session_id in session_dict:
        session_dict[session_id] = sorted(session_dict[session_id], key=lambda x: x['created_at'])
    
    # Format the output as a list of dictionaries with session_id as the key
    result = [
        {
            'session_id': session_id,
            'messages': session_dict[session_id]
        }
        for session_id in sorted(session_dict.keys(), key=lambda x: session_dict[x][-1]['created_at'], reverse=True)
    ]
    
    # Return the formatted chat history
    return result


def get_history_user(db: Session, user: str):
    logger.info('get_history_user!')
    
    # Query the UserSession table filtered by the user and ordered by session_id and created_at
    sessions = db.query(UserSession).filter(UserSession.nbkid == user).order_by(UserSession.session_id, UserSession.created_at).all()
    
    # Group the sessions by session_id
    session_dict = defaultdict(list)
    for session in sessions:
        session_dict[session.session_id].append(session)
    
    # Sort the session groups by their latest message timestamp in descending order
    sorted_sessions = sorted(session_dict.items(), key=lambda x: x[1][-1].created_at, reverse=True)
    
    # Return the grouped and sorted chat history
    return sorted_sessions
    
    
from sqlalchemy.orm import Session
from sqlalchemy import asc, desc
from your_logging_module import logger  # Import your logger

# Assuming @log_performance is a custom decorator you've implemented
# @log_performance
def get_history_user(db: Session, user):
    logger.info('get_history_user!')
    
    # Query to get sessions of the user sorted by creation date (latest first)
    sessions = db.query(UserSession).filter(UserSession.nbkid == user).order_by(desc(UserSession.created_at)).all()
    
    # Create a dictionary to store the chat history grouped by session
    history = {}
    
    for session in sessions:
        # Query to get chats in each session sorted by creation date (ascending)
        chats = db.query(Chat).filter(Chat.session_id == session.id).order_by(asc(Chat.created_at)).all()
        history[session.id] = {
            'session_info': session,
            'chats': chats
        }
    
    return history







from langchain.llms import HuggingFacePipeline
from langchain.prompts import PromptTemplate

def create_prompt():
    template = """You are a question answering Large Language Model.
    Answer the following question.
    USER: {question}
    ASSISTANT:"""
    prompt = PromptTemplate.from_template(template)
    return prompt

def create_llm_pipeline(model_id):
    llm_pipeline = HuggingFacePipeline.from_model_id(
    model_id=model_id,
    task="text-generation",
    device=0,  # -1 for CPU
    batch_size=1,  
    model_kwargs={"do_sample": True, "max_length": 128},
    )
    return llm_pipeline

def create_chain(prompt, llm_pipeline):
    chain = prompt | llm_pipeline.bind(stop=["USER:"])
    return chain

if __name__ == "__main__":
    model_id = "togethercomputer/RedPajama-INCITE-Instruct-3B-v1"
    #model_id = "mistralai/Mistral-7B-v0.1"
    chain = create_chain(create_prompt(), create_llm_pipeline(model_id=model_id))

    keep_running = True
    while keep_running:
        print("input something...")
        question = input()
        if question != "exit":
            print(chain.invoke({"question": question}))
        else:
            keep_running = False
